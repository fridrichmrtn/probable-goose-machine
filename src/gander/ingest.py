from __future__ import annotations

import asyncio
import os
import re
import time
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import docx
import fitz  # type: ignore[import-untyped]
import httpx
import pdfplumber
import pypdf

from gander import obs
from gander.config import env_float, env_int
from gander.errors import StageFailure, stage_boundary
from gander.llm import get_client
from gander.sections import NORMALIZED_SECTION_NAMES, SECTION_NAMES, normalize_section_name

SCANNED_MSG = "This appears to be a scanned PDF. Text-based PDFs and DOCX are required."
UNKNOWN_MSG = "Unable to read this file. Please upload a valid PDF or DOCX."
DOC_MSG = "Legacy .doc is not supported. Please convert to PDF or DOCX and re-upload."
CORRUPT_MSG = "Could not read this file. It may be corrupt or password-protected."
EMPTY_MSG = "This file appears to be empty or too short to be a CV."
LOW_EVIDENCE_MSG = (
    "We couldn't find the experience, education, or skills we look for in a CV. "
    "If this is a CV, check that the text is selectable (not a scanned image) "
    "and that sections like Experience or Education are clearly labelled, then try again."
)
VISION_BUDGET_MSG = (
    "This PDF has too many pages to transcribe reliably. Try a shorter PDF, "
    "or upload a text-based PDF where text is selectable."
)
VISION_BUDGET_FALLBACK_NOTICE = "Vision skipped: PDF over budget; used text extraction."

MIN_TEXT_CHARS = 100
# The length gate exists ONLY to keep long paragraph sentences that contain a
# section keyword (e.g. "He learned several languages…") out of the vocab-match
# branch. 40 chars covers every alias plus typical decorations (trailing colon,
# single trailing word). It does NOT apply to the all-caps branch: real CVs
# sometimes use long all-caps section headers, especially in Czech where
# diacritics make labels longer (e.g. "PRACOVNÍ ZKUŠENOSTI A PROFESNÍ KARIÉRA").
_MAX_HEADER_CHARS = 40
_INGEST_MODES = {"vision", "text"}
_PDF_INGEST_MODES = {"vision", "text"}
_DOCX_INGEST_MODES = {"llm", "text"}
_DEFAULT_VISION_DPI = 160
_DEFAULT_VISION_MAX_PAGES = 8
# A4 at 160 dpi is roughly 2.2 Mpx; this allows larger/slides-style pages
# without letting a single rendered page dominate memory or provider spend.
_DEFAULT_VISION_MAX_PIXELS_PER_PAGE = 6_000_000
# Aggregate rendered PNG budget. This protects provider payload and cost even
# when each individual page is within the pixel budget.
_DEFAULT_VISION_MAX_TOTAL_IMAGE_BYTES = 48_000_000
_DEFAULT_VISION_MAX_PDF_BYTES = 10_000_000
_DEFAULT_VISION_CONCURRENCY = 4
_MIN_DOCX_SOURCE_OVERLAP = 0.70
_PDF_MAGIC = b"%PDF"
# DOCX is a ZIP container; PK\x03\x04 is the ZIP local-file-header signature.
_DOCX_MAGIC = b"PK\x03\x04"
_DEFAULT_MAX_INPUT_CHARS = 50_000

_SIDEBAR_TOKENS = (
    "kontakt",
    "contact",
    "nejčastější dovednosti",
    "dovednosti",
    "skills",
    "jazyky",
    "languages",
    "certifikace",
    "certifications",
    "honors-awards",
)
_BODY_TOKENS = (
    "pracovní zkušenosti",
    "professional experience",
    "work experience",
    "experience",
)


class _IngestLLMReject(Exception):
    def __init__(self, reason: str) -> None:
        self.reason = reason
        super().__init__(reason)


class _VisionBudgetExceeded(Exception):
    def __init__(self, reason: str) -> None:
        self.reason = reason
        super().__init__(reason)


@dataclass(frozen=True)
class _RenderedPdf:
    pages: list[bytes]
    total_image_bytes: int
    max_page_pixels: int


def _load_prompt(name: str) -> str:
    return (Path(__file__).with_name("prompts") / name).read_text(encoding="utf-8")


def _check_magic_bytes(file_bytes: bytes, suffix: str) -> bool:
    """Return True when the file's leading bytes match its claimed suffix."""
    if suffix == ".pdf":
        return file_bytes[:4] == _PDF_MAGIC
    if suffix == ".docx":
        return file_bytes[:4] == _DOCX_MAGIC
    return True


def _max_input_chars() -> int:
    return env_int("GANDER_MAX_INPUT_CHARS", _DEFAULT_MAX_INPUT_CHARS, max_value=200_000)


def _ingest_mode() -> str:
    mode = os.environ.get("GANDER_INGEST_MODE", "vision").strip().lower()
    if mode not in _INGEST_MODES:
        expected = ", ".join(sorted(_INGEST_MODES))
        raise ValueError(f"Unknown GANDER_INGEST_MODE={mode!r}; expected one of {expected}")
    return mode


def _read_mode_env(env_name: str, allowed: set[str]) -> str:
    mode = os.environ.get(env_name, "").strip().lower()
    if mode not in allowed:
        expected = ", ".join(sorted(allowed))
        raise ValueError(f"Unknown {env_name}={mode!r}; expected one of {expected}")
    return mode


def _pdf_ingest_mode() -> str:
    if "GANDER_PDF_INGEST_MODE" in os.environ:
        return _read_mode_env("GANDER_PDF_INGEST_MODE", _PDF_INGEST_MODES)
    return _ingest_mode()


def _docx_ingest_mode() -> str:
    if "GANDER_DOCX_INGEST_MODE" in os.environ:
        return _read_mode_env("GANDER_DOCX_INGEST_MODE", _DOCX_INGEST_MODES)
    if "GANDER_INGEST_MODE" in os.environ:
        return "llm" if _ingest_mode() == "vision" else "text"
    return "text"


def _vision_dpi() -> int:
    raw = os.environ.get("GANDER_VISION_DPI")
    if raw is None:
        return _DEFAULT_VISION_DPI
    try:
        return min(300, max(120, int(raw)))
    except ValueError:
        return _DEFAULT_VISION_DPI


def _vision_max_pages() -> int:
    return env_int("GANDER_VISION_MAX_PAGES", _DEFAULT_VISION_MAX_PAGES)


def _vision_max_pixels_per_page() -> int:
    return env_int(
        "GANDER_VISION_MAX_PIXELS_PER_PAGE",
        _DEFAULT_VISION_MAX_PIXELS_PER_PAGE,
    )


def _vision_max_total_image_bytes() -> int:
    return env_int(
        "GANDER_VISION_MAX_TOTAL_IMAGE_BYTES",
        _DEFAULT_VISION_MAX_TOTAL_IMAGE_BYTES,
    )


def _vision_max_pdf_bytes() -> int:
    return env_int("GANDER_VISION_MAX_PDF_BYTES", _DEFAULT_VISION_MAX_PDF_BYTES)


def _vision_concurrency() -> int:
    return env_int("GANDER_VISION_CONCURRENCY", _DEFAULT_VISION_CONCURRENCY, max_value=8)


def _vision_timeout_s() -> float:
    return env_float("GANDER_VISION_TIMEOUT_S", 120.0)


def _is_all_caps_header(s: str) -> bool:
    """Unicode-aware all-caps section heuristic.

    Accepts a line that is at least 7 chars long, starts with an uppercase
    letter, and contains only uppercase letters plus the decorations seen on
    real CV headers (space, ampersand, slash). Unicode-aware so Czech
    diacritics like Í/Š/É count as uppercase letters — the previous ASCII-only
    regex silently dropped every all-caps Czech header.
    """
    if len(s) < 7 or not s[0].isupper():
        return False
    has_letter = False
    for ch in s:
        if ch.isalpha():
            if not ch.isupper():
                return False
            has_letter = True
        elif ch not in " &/":
            return False
    return has_letter


def _section_repair_aliases() -> tuple[str, ...]:
    aliases: set[str] = set()
    for name in SECTION_NAMES:
        aliases.add(name[:1].upper() + name[1:])
        aliases.add(" ".join(part[:1].upper() + part[1:] for part in name.split(" ")))
        aliases.add(name.upper())
    return tuple(sorted(aliases, key=len, reverse=True))


_INLINE_SECTION_RE = re.compile(
    r"(?<![#\w])("
    + "|".join(re.escape(alias) for alias in _section_repair_aliases())
    + r")(?=\s|:|$)"
)
_TOKEN_RE = re.compile(r"[0-9A-Za-zÀ-ž]+", re.UNICODE)


async def extract_text(file_bytes: bytes, filename: str) -> str | StageFailure:
    """Extract plain text from a CV file (PDF or DOCX).

    Returns the text on success, or a StageFailure with a user-facing message on
    controlled failure paths (unknown suffix, legacy .doc, scanned PDF, corrupt
    file). Wrapped in stage_boundary as defense-in-depth for unexpected errors.
    """
    t0 = time.perf_counter()
    suffix = Path(filename).suffix.lower()
    obs.emit("ingest", "start", filename_suffix=suffix, size_bytes=len(file_bytes))

    def _ms() -> int:
        return int((time.perf_counter() - t0) * 1000)

    async with stage_boundary("ingest") as cm:
        if suffix == ".doc":
            obs.emit("ingest", "rejected", reason="doc_legacy", duration_ms=_ms())
            return StageFailure(stage="ingest", user_message=DOC_MSG)

        if not _check_magic_bytes(file_bytes, suffix):
            obs.emit(
                "ingest",
                "rejected",
                reason="wrong_magic_bytes",
                suffix=suffix,
                size_bytes=len(file_bytes),
                duration_ms=_ms(),
            )
            return StageFailure(
                stage="ingest",
                user_message=CORRUPT_MSG,
                debug_detail=f"magic-byte mismatch for {suffix}: {len(file_bytes)} bytes",
            )

        mode: str
        if suffix == ".pdf":
            mode = _pdf_ingest_mode()
            try:
                text = await _extract_pdf(file_bytes, mode=mode)
            except _VisionBudgetExceeded as exc:
                obs.emit(
                    "ingest",
                    "rejected",
                    reason="vision_budget_exceeded",
                    detail=exc.reason,
                    duration_ms=_ms(),
                )
                return StageFailure(
                    stage="ingest",
                    user_message=VISION_BUDGET_MSG,
                    debug_detail=exc.reason,
                )
            except Exception as exc:
                obs.emit("ingest", "rejected", reason="corrupt", duration_ms=_ms())
                return StageFailure(
                    stage="ingest",
                    user_message=CORRUPT_MSG,
                    debug_detail=f"{type(exc).__name__}: {len(file_bytes)} bytes",
                )
            if len(text.strip()) < MIN_TEXT_CHARS:
                obs.emit("ingest", "rejected", reason="scanned", duration_ms=_ms())
                return StageFailure(stage="ingest", user_message=SCANNED_MSG)
        elif suffix == ".docx":
            mode = _docx_ingest_mode()
            try:
                text = await _extract_docx(file_bytes, mode=mode)
            except Exception as exc:
                obs.emit("ingest", "rejected", reason="corrupt", duration_ms=_ms())
                return StageFailure(
                    stage="ingest",
                    user_message=CORRUPT_MSG,
                    debug_detail=f"{type(exc).__name__}: {len(file_bytes)} bytes",
                )
            if len(text.strip()) < MIN_TEXT_CHARS:
                obs.emit("ingest", "rejected", reason="empty_docx", duration_ms=_ms())
                return StageFailure(stage="ingest", user_message=EMPTY_MSG)
        else:
            obs.emit("ingest", "rejected", reason="unknown_suffix", duration_ms=_ms())
            return StageFailure(stage="ingest", user_message=UNKNOWN_MSG)

        repaired = _repair_inline_section_breaks(text)
        annotated = _annotate_sections(repaired)
        max_chars = _max_input_chars()
        if len(annotated) > max_chars:
            obs.emit(
                "ingest",
                "input_truncated",
                original_chars=len(annotated),
                max_chars=max_chars,
            )
            annotated = annotated[:max_chars]
        obs.emit(
            "ingest",
            "done",
            chars=len(annotated),
            suffix=suffix,
            duration_ms=_ms(),
            mode=mode,
        )
        return annotated

    assert cm.failure is not None  # stage_boundary swallowed an unexpected exception
    return cm.failure


async def _extract_pdf(file_bytes: bytes, *, mode: str | None = None) -> str:
    resolved_mode = _pdf_ingest_mode() if mode is None else mode
    if resolved_mode == "text":
        return await asyncio.to_thread(_extract_pdf_text, file_bytes)

    try:
        return await _extract_pdf_vlm(file_bytes)
    except _VisionBudgetExceeded as exc:
        obs.emit("ingest", "ingest_llm_fallback", file_type="pdf", reason=exc.reason)
        text = await asyncio.to_thread(_extract_pdf_text, file_bytes)
        if len(text.strip()) >= MIN_TEXT_CHARS:
            obs.emit(
                "ingest",
                "vision_budget_fallback_degraded",
                file_type="pdf",
                reason=exc.reason,
                notice=VISION_BUDGET_FALLBACK_NOTICE,
            )
            return text
        raise
    except _IngestLLMReject as exc:
        obs.emit("ingest", "ingest_llm_fallback", file_type="pdf", reason=exc.reason)
    except (httpx.HTTPError, RuntimeError, TimeoutError, ValueError) as exc:
        obs.emit(
            "ingest",
            "ingest_llm_fallback",
            file_type="pdf",
            reason="api_error",
            exc_type=type(exc).__name__,
        )
    return await asyncio.to_thread(_extract_pdf_text, file_bytes)


async def _extract_docx(file_bytes: bytes, *, mode: str | None = None) -> str:
    resolved_mode = _docx_ingest_mode() if mode is None else mode
    deterministic = await asyncio.to_thread(_extract_docx_text, file_bytes)
    if resolved_mode == "text":
        return deterministic
    if len(deterministic.strip()) < MIN_TEXT_CHARS:
        return deterministic

    try:
        candidate = await _normalize_docx_with_llm(deterministic)
        candidate = _repair_inline_section_breaks(candidate)
        _validate_docx_llm_candidate(deterministic, candidate)
        obs.emit(
            "ingest",
            "docx_llm_done",
            source_chars=len(deterministic),
            chars=len(candidate),
            source_overlap=_source_overlap(deterministic, candidate),
        )
        return candidate
    except Exception as exc:
        reason = exc.reason if isinstance(exc, _IngestLLMReject) else "api_error"
        obs.emit(
            "ingest",
            "ingest_llm_fallback",
            file_type="docx",
            reason=reason,
            exc_type=type(exc).__name__,
        )
        return deterministic


def _extract_pdf_text(file_bytes: bytes) -> str:
    reader = pypdf.PdfReader(BytesIO(file_bytes))
    pypdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    pypdf_chars = len(pypdf_text.strip())

    if pypdf_chars >= MIN_TEXT_CHARS:
        obs.emit(
            "ingest",
            "pdf_pass",
            pypdf_chars=pypdf_chars,
            pdfplumber_chars=0,
            used="pypdf",
        )
        return pypdf_text

    # pypdf produced something parseable but text-poor. If pdfplumber blows up,
    # don't reclassify the file as corrupt — fall back to pypdf's output and
    # let the caller's <100-char check yield SCANNED_MSG.
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as plumber_pdf:
            plumber_text = "\n".join(page.extract_text() or "" for page in plumber_pdf.pages)
        plumber_chars = len(plumber_text.strip())
    except Exception as exc:
        obs.emit(
            "ingest",
            "pdfplumber_fallback_failed",
            pypdf_chars=pypdf_chars,
            exc_type=type(exc).__name__,
        )
        return pypdf_text

    used = "pdfplumber" if plumber_chars > pypdf_chars else "pypdf"
    obs.emit(
        "ingest",
        "pdf_pass",
        pypdf_chars=pypdf_chars,
        pdfplumber_chars=plumber_chars,
        used=used,
    )
    return plumber_text if used == "pdfplumber" else pypdf_text


def _extract_docx_text(file_bytes: bytes) -> str:
    document = docx.Document(BytesIO(file_bytes))
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _render_pdf_pages(file_bytes: bytes, dpi: int | None = None) -> list[bytes]:
    render_dpi = _vision_dpi() if dpi is None else dpi
    pages: list[bytes] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        for page in doc:
            pix = page.get_pixmap(dpi=render_dpi, alpha=False)
            pages.append(pix.tobytes("png"))
    finally:
        doc.close()
    return pages


def _render_pdf_pages_for_vision(file_bytes: bytes, dpi: int) -> _RenderedPdf:
    max_pages = _vision_max_pages()
    max_pixels = _vision_max_pixels_per_page()
    max_total_bytes = _vision_max_total_image_bytes()
    max_pdf_bytes = _vision_max_pdf_bytes()
    pages: list[bytes] = []
    total_image_bytes = 0
    max_page_pixels = 0
    if len(file_bytes) > max_pdf_bytes:
        obs.emit(
            "ingest",
            "vision_budget_rejected",
            reason="pdf_bytes",
            pdf_bytes=len(file_bytes),
            max_pdf_bytes=max_pdf_bytes,
        )
        raise _VisionBudgetExceeded(f"pdf_bytes={len(file_bytes)} max_pdf_bytes={max_pdf_bytes}")
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        page_count = doc.page_count
        if page_count > max_pages:
            obs.emit(
                "ingest",
                "vision_budget_rejected",
                reason="too_many_pages",
                page_count=page_count,
                max_pages=max_pages,
            )
            raise _VisionBudgetExceeded(f"page_count={page_count} max_pages={max_pages}")
        for page_index, page in enumerate(doc):
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            page_pixels = pix.width * pix.height
            max_page_pixels = max(max_page_pixels, page_pixels)
            if page_pixels > max_pixels:
                obs.emit(
                    "ingest",
                    "vision_budget_rejected",
                    reason="page_pixels",
                    page_index=page_index,
                    page_pixels=page_pixels,
                    max_pixels=max_pixels,
                )
                raise _VisionBudgetExceeded(
                    f"page_index={page_index} page_pixels={page_pixels} max_pixels={max_pixels}"
                )
            png = pix.tobytes("png")
            total_image_bytes += len(png)
            if total_image_bytes > max_total_bytes:
                obs.emit(
                    "ingest",
                    "vision_budget_rejected",
                    reason="total_image_bytes",
                    page_count=len(pages) + 1,
                    total_image_bytes=total_image_bytes,
                    max_total_image_bytes=max_total_bytes,
                )
                raise _VisionBudgetExceeded(
                    f"total_image_bytes={total_image_bytes} max_total_image_bytes={max_total_bytes}"
                )
            pages.append(png)
    finally:
        doc.close()
    return _RenderedPdf(
        pages=pages,
        total_image_bytes=total_image_bytes,
        max_page_pixels=max_page_pixels,
    )


async def _extract_pdf_vlm(file_bytes: bytes) -> str:
    dpi = _vision_dpi()
    rendered = await asyncio.to_thread(_render_pdf_pages_for_vision, file_bytes, dpi=dpi)
    pages = rendered.pages
    if not pages:
        raise _IngestLLMReject("page_render_failed")

    prompt = _load_prompt("ingest_vlm.md")
    client = get_client()
    obs.emit(
        "ingest",
        "ingest_vlm_start",
        page_count=len(pages),
        dpi=dpi,
        size_bytes=len(file_bytes),
        total_image_bytes=rendered.total_image_bytes,
        max_page_pixels=rendered.max_page_pixels,
        vision_pages_sent=len(pages),
    )

    sem = asyncio.Semaphore(_vision_concurrency())
    timeout_s = _vision_timeout_s()

    async def _transcribe(i: int, png: bytes) -> str:
        page_t0 = time.perf_counter()
        async with sem:
            page_text = await client.complete_vision_text(
                image_bytes=png,
                prompt=prompt,
                timeout_s=timeout_s,
                max_tokens=1500,
            )
        page_text = _strip_transcript_fences(page_text)
        if not page_text.strip():
            raise _IngestLLMReject("empty_output")
        obs.emit(
            "ingest",
            "ingest_vlm_page_done",
            page_index=i,
            chars=len(page_text),
            duration_ms=int((time.perf_counter() - page_t0) * 1000),
        )
        return page_text.strip()

    try:
        async with asyncio.TaskGroup() as tg:
            tasks = [tg.create_task(_transcribe(i, png)) for i, png in enumerate(pages)]
    except* (_IngestLLMReject, httpx.HTTPError, RuntimeError, TimeoutError, ValueError) as eg:
        # TaskGroup cancels still-running siblings before re-raising; surface the
        # first leaf exception so callers see the same single-exception contract
        # the pre-parallel serial loop used to deliver.
        leaf: BaseException = eg.exceptions[0]
        while isinstance(leaf, BaseExceptionGroup):
            leaf = leaf.exceptions[0]
        raise leaf from eg
    transcripts = [t.result() for t in tasks]

    transcript = "\n[PAGE_BREAK]\n".join(transcripts)
    transcript = _repair_inline_section_breaks(transcript)
    _validate_pdf_vlm_candidate(transcript)
    obs.emit("ingest", "ingest_vlm_done", chars=len(transcript), pages=len(transcripts))
    return transcript


async def _normalize_docx_with_llm(source_text: str) -> str:
    system = _load_prompt("ingest_docx.md")
    user = (
        "Normalize this deterministic DOCX extraction into a CV transcript.\n\n"
        "Return only the normalized transcript.\n\n"
        "SOURCE DOCX TEXT:\n"
        f"{source_text}"
    )
    obs.emit("ingest", "docx_llm_start", source_chars=len(source_text))
    text = await get_client().complete_text(
        system=system, user=user, model="cheap", temperature=0.0
    )
    return _strip_transcript_fences(text)


def _strip_transcript_fences(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```") and stripped.endswith("```"):
        lines = stripped.splitlines()
        if len(lines) >= 2:
            return "\n".join(lines[1:-1]).strip()
    return stripped


def _validate_pdf_vlm_candidate(candidate: str) -> None:
    stripped = candidate.strip()
    if len(stripped) < MIN_TEXT_CHARS:
        raise _IngestLLMReject("under_min_chars")
    if not _sidebar_body_order_ok(stripped):
        raise _IngestLLMReject("column_order_flipped")


def _validate_docx_llm_candidate(source: str, candidate: str) -> None:
    stripped = candidate.strip()
    if len(stripped) < MIN_TEXT_CHARS:
        raise _IngestLLMReject("under_min_chars")
    if len(stripped) < max(MIN_TEXT_CHARS, int(len(source.strip()) * 0.45)):
        raise _IngestLLMReject("under_min_chars")
    overlap = _source_overlap(source, candidate)
    if overlap < _MIN_DOCX_SOURCE_OVERLAP:
        raise _IngestLLMReject("low_source_overlap")
    if not _sidebar_body_order_ok(stripped):
        raise _IngestLLMReject("column_order_flipped")


def _source_overlap(source: str, candidate: str) -> float:
    source_tokens = {token for token in _tokens(source) if len(token) >= 3}
    if not source_tokens:
        return 1.0
    candidate_tokens = {token for token in _tokens(candidate) if len(token) >= 3}
    return len(source_tokens & candidate_tokens) / len(source_tokens)


def _tokens(text: str) -> list[str]:
    return [match.group(0).casefold() for match in _TOKEN_RE.finditer(text)]


def _norm_for_position(text: str) -> str:
    return " ".join(_tokens(text))


def _sidebar_body_order_ok(text: str) -> bool:
    norm_text = _norm_for_position(text)
    if not norm_text:
        return True

    sidebar_positions = [
        norm_text.find(_norm_for_position(token))
        for token in _SIDEBAR_TOKENS
        if norm_text.find(_norm_for_position(token)) >= 0
    ]
    body_positions = [
        norm_text.find(_norm_for_position(token))
        for token in _BODY_TOKENS
        if norm_text.find(_norm_for_position(token)) >= 0
    ]
    if not sidebar_positions or not body_positions:
        return True

    contact_positions = [
        norm_text.find(token) for token in ("kontakt", "contact") if norm_text.find(token) >= 0
    ]
    first_body = min(body_positions)
    if contact_positions and min(contact_positions) > first_body:
        return False
    if len(sidebar_positions) >= 3 and min(sidebar_positions) < first_body:
        return max(sidebar_positions) < first_body
    return True


def _repair_inline_section_breaks(text: str) -> str:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    return "\n".join(_repair_inline_section_line(line) for line in lines)


def _repair_inline_section_line(line: str) -> str:
    stripped = line.strip()
    if not stripped or stripped.startswith("##"):
        return line
    matches = list(_INLINE_SECTION_RE.finditer(line))
    if not matches:
        return line

    parts: list[str] = []
    cursor = 0
    for match in matches:
        prefix = line[cursor : match.start()].strip()
        if prefix:
            parts.append(prefix)
        heading = match.group(1).strip()
        cursor = match.end()
        if cursor < len(line) and line[cursor] == ":":
            cursor += 1
        parts.append(heading)
    suffix = line[cursor:].strip()
    if suffix:
        parts.append(suffix)
    return "\n".join(parts)


def _looks_like_section_header(line: str) -> bool:
    raw = line.strip()
    if raw.startswith("##"):
        return False
    stripped = raw.rstrip(":").strip()
    if not stripped:
        return False
    # All-caps branch: trusted enough to bypass the length gate. The strict
    # shape (uppercase letters + space/&// only) already rules out paragraph
    # sentences. Long Czech all-caps labels live here.
    if _is_all_caps_header(stripped):
        return True
    # Vocab-match branch: the length gate applies here because this branch is
    # the one prone to false positives — paragraph sentences containing a
    # section keyword normalise to a vocab miss only because the surrounding
    # words push them past the threshold. Without the gate, "He learned
    # several languages…" risks promotion if it ever shortens.
    if len(stripped) > _MAX_HEADER_CHARS:
        return False
    return normalize_section_name(stripped) in NORMALIZED_SECTION_NAMES


def _annotate_sections(text: str) -> str:
    lines = text.split("\n")
    out: list[str] = []
    for i, line in enumerate(lines):
        if _looks_like_section_header(line):
            prev = out[-1].strip() if out else ""
            already_annotated = prev.startswith("##") and prev.lstrip("#").strip().lower() == (
                line.strip().lower()
            )
            previous_source = lines[i - 1].strip() if i > 0 else ""
            already_in_source = previous_source.startswith("##") and previous_source.lstrip(
                "#"
            ).strip().lower() == (line.strip().lower())
            if not already_annotated and not already_in_source:
                out.append(f"## {line.strip()}")
        out.append(line)
    return "\n".join(out)
