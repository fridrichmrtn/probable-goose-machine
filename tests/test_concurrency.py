"""Concurrency hygiene (P0.3): shared LLM client, parser offload to threads."""

from __future__ import annotations

import asyncio
from io import BytesIO
from typing import Any

import pytest

from gander import ingest
from gander.llm import LLMClient, get_client


@pytest.fixture(autouse=True)
def _api_key(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("OPENROUTER_API_KEY", "test-stub")


@pytest.mark.fast
def test_get_client_returns_same_instance() -> None:
    assert get_client() is get_client()


@pytest.mark.fast
def test_get_client_cache_clear_produces_new_instance() -> None:
    a = get_client()
    get_client.cache_clear()
    assert a is not get_client()


@pytest.mark.fast
async def test_pdf_text_extraction_runs_in_thread(monkeypatch: pytest.MonkeyPatch) -> None:
    reportlab_canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    buf = BytesIO()
    c = reportlab_canvas.Canvas(buf)
    y = 760
    for line in (
        "Work Experience",
        "Built a fraud-detection service using PyTorch and Kafka stream processing.",
        "Owned the on-call rotation across two production squads for eighteen months.",
    ):
        c.drawString(72, y, line)
        y -= 24
    c.showPage()
    c.save()
    pdf = buf.getvalue()

    offloaded: list[str] = []
    real_to_thread = asyncio.to_thread

    async def spy_to_thread(fn: Any, /, *args: Any, **kwargs: Any) -> Any:
        offloaded.append(fn.__name__)
        return await real_to_thread(fn, *args, **kwargs)

    monkeypatch.setattr(ingest.asyncio, "to_thread", spy_to_thread)

    text = await ingest._extract_pdf(pdf, mode="text")

    assert "_extract_pdf_text" in offloaded
    assert "fraud-detection" in text


@pytest.mark.fast
async def test_pdf_vision_render_runs_in_thread(monkeypatch: pytest.MonkeyPatch) -> None:
    reportlab_canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    buf = BytesIO()
    c = reportlab_canvas.Canvas(buf)
    c.drawString(72, 760, "Work Experience")
    c.showPage()
    c.save()
    pdf = buf.getvalue()

    async def _fake_vlm(
        self: LLMClient,
        *,
        image_bytes: bytes,
        prompt: str,
        mime_type: str = "image/png",
        timeout_s: float = 120.0,
        max_tokens: int | None = None,
    ) -> str:
        return (
            "Work Experience\n"
            "Built a fraud-detection service using PyTorch and Kafka stream processing. "
            "Owned the on-call rotation across two production squads for eighteen months."
        )

    monkeypatch.setattr(LLMClient, "complete_vision_text", _fake_vlm)

    offloaded: list[str] = []
    real_to_thread = asyncio.to_thread

    async def spy_to_thread(fn: Any, /, *args: Any, **kwargs: Any) -> Any:
        offloaded.append(fn.__name__)
        return await real_to_thread(fn, *args, **kwargs)

    monkeypatch.setattr(ingest.asyncio, "to_thread", spy_to_thread)

    text = await ingest._extract_pdf(pdf, mode="vision")

    assert "_render_pdf_pages_for_vision" in offloaded
    assert "fraud-detection" in text


@pytest.mark.fast
async def test_docx_text_extraction_runs_in_thread(monkeypatch: pytest.MonkeyPatch) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Work Experience")
    document.add_paragraph("Shipped a production recommender system used by two million users.")
    buf = BytesIO()
    document.save(buf)

    offloaded: list[str] = []
    real_to_thread = asyncio.to_thread

    async def spy_to_thread(fn: Any, /, *args: Any, **kwargs: Any) -> Any:
        offloaded.append(fn.__name__)
        return await real_to_thread(fn, *args, **kwargs)

    monkeypatch.setattr(ingest.asyncio, "to_thread", spy_to_thread)

    text = await ingest._extract_docx(buf.getvalue(), mode="text")

    assert "_extract_docx_text" in offloaded
    assert "recommender system" in text
