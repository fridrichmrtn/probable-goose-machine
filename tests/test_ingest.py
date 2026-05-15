from __future__ import annotations

import asyncio
import unicodedata
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest

from gander import ingest
from gander.errors import StageFailure
from gander.ingest import (
    CORRUPT_MSG,
    DOC_MSG,
    EMPTY_MSG,
    SCANNED_MSG,
    UNKNOWN_MSG,
    _annotate_sections,
    _repair_inline_section_breaks,
    extract_text,
)
from gander.llm import LLMClient
from gander.obs import subscribe


@pytest.fixture(autouse=True)
def _deterministic_ingest(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("GANDER_INGEST_MODE", "text")


def _docx_bytes(paragraphs: list[str]) -> bytes:
    import docx as _docx

    document = _docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


def _pdf_bytes(pages: list[list[str]]) -> bytes:
    reportlab_canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    buf = BytesIO()
    c = reportlab_canvas.Canvas(buf)
    for lines in pages:
        y = 760
        for line in lines:
            c.drawString(72, y, line)
            y -= 24
        c.showPage()
    c.save()
    return buf.getvalue()


@pytest.mark.fast
async def test_unknown_extension_returns_format_failure() -> None:
    result = await extract_text(b"hello", "notes.txt")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == UNKNOWN_MSG


@pytest.mark.fast
async def test_doc_extension_returns_conversion_hint() -> None:
    result = await extract_text(b"\xd0\xcf\x11\xe0fake-ole-bytes", "cv.doc")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == DOC_MSG


@pytest.mark.fast
async def test_corrupt_pdf_returns_corrupt_failure() -> None:
    result = await extract_text(b"%PDF-not-a-real-pdf", "cv.pdf")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == CORRUPT_MSG
    assert result.user_message != SCANNED_MSG


@pytest.mark.fast
async def test_corrupt_pdf_debug_detail_does_not_leak_content() -> None:
    payload = b"%PDF-not-a-real-pdf"
    result = await extract_text(payload, "cv.pdf")
    assert isinstance(result, StageFailure)
    assert result.debug_detail is not None
    assert f"{len(payload)} bytes" in result.debug_detail
    assert "not-a-real-pdf" not in result.debug_detail


@pytest.mark.fast
async def test_corrupt_docx_returns_corrupt_failure() -> None:
    result = await extract_text(b"PK\x03\x04junk-not-a-real-docx", "cv.docx")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == CORRUPT_MSG


@pytest.mark.fast
async def test_empty_docx_returns_empty_failure() -> None:
    import docx as _docx

    document = _docx.Document()
    buf = BytesIO()
    document.save(buf)
    result = await extract_text(buf.getvalue(), "empty.docx")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == EMPTY_MSG


@pytest.mark.fast
async def test_tiny_docx_returns_empty_failure() -> None:
    import docx as _docx

    document = _docx.Document()
    document.add_paragraph("Hi")
    buf = BytesIO()
    document.save(buf)
    result = await extract_text(buf.getvalue(), "tiny.docx")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == EMPTY_MSG


@pytest.mark.fast
async def test_pdfplumber_failure_falls_back_to_scanned_not_corrupt(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """When pypdf parses a valid PDF but extracts <100 chars, and pdfplumber
    then raises during fallback, classify as SCANNED (text-poor), not CORRUPT.
    """
    reportlab_canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    buf = BytesIO()
    c = reportlab_canvas.Canvas(buf)
    c.drawString(100, 100, "hi")  # tiny but pypdf-parseable text
    c.showPage()
    c.save()
    pdf_bytes = buf.getvalue()

    def _boom(*_args: object, **_kwargs: object) -> object:
        raise RuntimeError("synthetic pdfplumber failure")

    monkeypatch.setattr(ingest.pdfplumber, "open", _boom)

    result = await extract_text(pdf_bytes, "tiny.pdf")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == SCANNED_MSG


@pytest.mark.fast
def test_render_pdf_pages_returns_pngs() -> None:
    pdf_bytes = _pdf_bytes(
        [
            ["Summary", "Data Scientist with Python and SQL."],
            ["Experience", "Built forecasting systems for retail teams."],
        ]
    )

    pages = ingest._render_pdf_pages(pdf_bytes, dpi=120)

    assert len(pages) == 2
    assert all(page.startswith(b"\x89PNG\r\n\x1a\n") for page in pages)


@pytest.mark.fast
async def test_pdf_vlm_ingest_renders_pages_and_joins_transcripts(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("GANDER_INGEST_MODE", "vision")
    monkeypatch.setenv("MINIMAX_API_KEY", "test-stub")
    pdf_bytes = _pdf_bytes(
        [
            ["Summary", "Data Scientist with 5 years building churn models."],
            ["Experience", "Led customer churn model for a synthetic retail pilot."],
        ]
    )
    seen_images: list[bytes] = []

    async def _fake_vlm(
        self: LLMClient,
        *,
        image_bytes: bytes,
        prompt: str,
        mime_type: str = "image/png",
        timeout_s: float = 120.0,
        max_tokens: int | None = None,
    ) -> str:
        assert "Transcribe this CV page verbatim" in prompt
        assert mime_type == "image/png"
        assert timeout_s == 120.0
        assert max_tokens == 1500
        seen_images.append(image_bytes)
        if len(seen_images) == 1:
            return (
                "Summary Data Scientist with 5 years building churn models for retail teams. "
                "Skills Python SQL LightGBM model monitoring dashboards."
            )
        return (
            "Experience Led customer churn model for a synthetic retail pilot reducing "
            "cancellations by 11 percent. Education CVUT FIT Prague MSc Informatics."
        )

    monkeypatch.setattr(LLMClient, "complete_vision_text", _fake_vlm)

    result = await extract_text(pdf_bytes, "cv.pdf")

    assert isinstance(result, str)
    assert len(seen_images) == 2
    assert all(image.startswith(b"\x89PNG\r\n\x1a\n") for image in seen_images)
    assert "[PAGE_BREAK]" in result
    assert "Data Scientist with 5 years building churn models" in result
    assert "Led customer churn model" in result
    assert "## Summary" in result


@pytest.mark.fast
async def test_pdf_vlm_parallel_preserves_page_order_and_bounds_concurrency(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """VLM page transcription runs concurrently (≤4 in flight) and the joined
    transcript preserves original page order. Regression guard against silently
    re-serializing the gather loop."""
    monkeypatch.setenv("GANDER_INGEST_MODE", "vision")
    monkeypatch.setenv("MINIMAX_API_KEY", "test-stub")
    pdf_bytes = _pdf_bytes(
        [["Summary", f"Page {i} content text for parallel ingest test."] for i in range(6)]
    )

    page_pngs = ingest._render_pdf_pages(pdf_bytes)
    png_to_index = {png: i for i, png in enumerate(page_pngs)}
    peak: dict[str, int] = {"in_flight": 0, "max": 0}
    captured_kwargs: list[dict[str, Any]] = []

    async def _fake_vlm(
        self: LLMClient,
        *,
        image_bytes: bytes,
        prompt: str,
        mime_type: str = "image/png",
        timeout_s: float = 120.0,
        max_tokens: int | None = None,
    ) -> str:
        captured_kwargs.append({"max_tokens": max_tokens})
        peak["in_flight"] += 1
        peak["max"] = max(peak["max"], peak["in_flight"])
        try:
            await asyncio.sleep(0.02)
            idx = png_to_index[image_bytes]
            return (
                f"Summary Page {idx} parallel transcript with enough chars to pass "
                "the minimum text gate for ingest."
            )
        finally:
            peak["in_flight"] -= 1

    monkeypatch.setattr(LLMClient, "complete_vision_text", _fake_vlm)

    events: list[dict[str, Any]] = []
    with subscribe(events.append):
        result = await extract_text(pdf_bytes, "cv.pdf")

    assert isinstance(result, str)
    pages = result.split("[PAGE_BREAK]")
    assert len(pages) == 6
    for i, page in enumerate(pages):
        assert f"Page {i} parallel transcript" in page, (
            f"page {i} out of order or missing: {page!r}"
        )

    assert peak["max"] <= 4, f"semaphore bound violated: {peak['max']}"
    assert peak["max"] >= 2, (
        f"expected concurrent overlap; got max in-flight {peak['max']} — "
        "did someone re-serialize the gather?"
    )

    assert all(c["max_tokens"] == 1500 for c in captured_kwargs)

    page_done = [e for e in events if e["event"] == "ingest_vlm_page_done"]
    assert len(page_done) == 6
    assert {e["page_index"] for e in page_done} == set(range(6))


@pytest.mark.fast
async def test_pdf_vlm_failure_falls_back_to_deterministic_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("GANDER_INGEST_MODE", "vision")
    monkeypatch.setenv("MINIMAX_API_KEY", "test-stub")
    deterministic_phrase = "Deterministic fallback phrase for a text PDF with enough content"
    pdf_bytes = _pdf_bytes(
        [
            [
                "Summary",
                deterministic_phrase,
                "Experience includes Python SQL data quality monitoring and forecasting.",
                "Education CVUT FIT Prague MSc Informatics completed in 2021.",
            ]
        ]
    )

    async def _raise_vlm(self: LLMClient, **_kwargs: object) -> str:
        raise RuntimeError("synthetic VLM outage")

    monkeypatch.setattr(LLMClient, "complete_vision_text", _raise_vlm)

    result = await extract_text(pdf_bytes, "cv.pdf")

    assert isinstance(result, str)
    assert deterministic_phrase in result


@pytest.mark.fast
async def test_docx_text_llm_normalizes_transcript(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("GANDER_INGEST_MODE", "vision")
    monkeypatch.setenv("MINIMAX_API_KEY", "test-stub")
    source = [
        "Summary Data Scientist with Python SQL LightGBM monitoring and dashboards.",
        "Pracovní zkušenosti Vedla model odchodu zákazníků pro Česko v října 2024.",
        "Vzdělání ČVUT FIT Datová věda září 2019 června 2021.",
    ]

    async def _fake_text(
        self: LLMClient,
        *,
        system: str,
        user: str,
        model: str = "cheap",
        temperature: float = 0.0,
    ) -> str:
        assert "normalize deterministic DOCX text" in system
        assert "SOURCE DOCX TEXT" in user
        assert model == "cheap"
        assert temperature == 0.0
        return "\n".join(source)

    monkeypatch.setattr(LLMClient, "complete_text", _fake_text)

    result = await extract_text(_docx_bytes(source), "cv.docx")

    assert isinstance(result, str)
    assert "Pracovní zkušenosti" in result
    assert "října" in result
    assert "## Vzdělání" in result


@pytest.mark.fast
async def test_docx_text_llm_low_overlap_falls_back_to_deterministic_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("GANDER_INGEST_MODE", "vision")
    monkeypatch.setenv("MINIMAX_API_KEY", "test-stub")
    source = [
        "Summary Senior Data Scientist with Python SQL LightGBM monitoring dashboards.",
        "Experience Owned fraud model lifecycle for synthetic finance platform.",
        "Education CVUT FIT Prague MSc Informatics 2021.",
    ]

    async def _paraphrase(self: LLMClient, **_kwargs: object) -> str:
        return "A strong candidate with many useful abilities and a good background."

    monkeypatch.setattr(LLMClient, "complete_text", _paraphrase)

    result = await extract_text(_docx_bytes(source), "cv.docx")

    assert isinstance(result, str)
    assert "Owned fraud model lifecycle for synthetic finance platform" in result
    assert "A strong candidate" not in result


@pytest.mark.fast
async def test_docx_text_llm_provider_error_falls_back_to_deterministic_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("GANDER_INGEST_MODE", "vision")
    monkeypatch.setenv("MINIMAX_API_KEY", "test-stub")
    source = [
        "Summary Senior Data Scientist with Python SQL LightGBM monitoring dashboards.",
        "Experience Owned fraud model lifecycle for synthetic finance platform.",
        "Education CVUT FIT Prague MSc Informatics 2021.",
    ]

    class ProviderError(Exception):
        pass

    async def _provider_error(self: LLMClient, **_kwargs: object) -> str:
        raise ProviderError("synthetic provider outage")

    monkeypatch.setattr(LLMClient, "complete_text", _provider_error)

    events: list[dict[str, Any]] = []
    with subscribe(events.append):
        result = await extract_text(_docx_bytes(source), "cv.docx")

    assert isinstance(result, str)
    assert "Owned fraud model lifecycle for synthetic finance platform" in result
    fallback = [e for e in events if e["event"] == "ingest_llm_fallback"]
    assert fallback
    assert fallback[0]["file_type"] == "docx"
    assert fallback[0]["reason"] == "api_error"
    assert fallback[0]["exc_type"] == "ProviderError"


@pytest.mark.fast
async def test_docx_fixture_vision_mode_preserves_role_and_sections(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("GANDER_INGEST_MODE", "vision")
    monkeypatch.setenv("MINIMAX_API_KEY", "test-stub")
    fixture = _FIXTURE_DIR / "01_junior_da_novotny.docx"

    async def _echo_source(self: LLMClient, *, user: str, **_kwargs: object) -> str:
        return user.split("SOURCE DOCX TEXT:\n", 1)[1]

    monkeypatch.setattr(LLMClient, "complete_text", _echo_source)

    result = await extract_text(fixture.read_bytes(), fixture.name)

    assert isinstance(result, str)
    assert "Junior Data Analyst" in result
    assert "## Summary" in result
    assert "## Experience" in result


@pytest.mark.fast
def test_docx_overlap_guard_rejects_paraphrase() -> None:
    source = (
        "Summary Senior Data Scientist with Python SQL LightGBM monitoring dashboards. "
        "Experience Owned fraud model lifecycle for synthetic finance platform. "
        "Education CVUT FIT Prague MSc Informatics 2021."
    )
    paraphrase = "A strong candidate with useful abilities and an academic background."

    with pytest.raises(Exception) as exc:
        ingest._validate_docx_llm_candidate(source, paraphrase)

    assert "under_min_chars" in str(exc.value) or "low_source_overlap" in str(exc.value)


@pytest.mark.fast
def test_repair_inline_section_breaks_for_collapsed_vlm_text() -> None:
    collapsed = (
        "Jana Testovací Pracovní zkušenosti Vedla model pro Česko v října 2024 "
        "Vzdělání ČVUT FIT Datová věda"
    )

    repaired = _repair_inline_section_breaks(collapsed)
    annotated = _annotate_sections(repaired)

    assert "Jana Testovací\nPracovní zkušenosti\nVedla model" in repaired
    assert "\nVzdělání\nČVUT FIT" in repaired
    assert "## Pracovní zkušenosti" in annotated
    assert "## Vzdělání" in annotated


@pytest.mark.fast
def test_annotate_inserts_header_for_all_caps_line() -> None:
    out = _annotate_sections("EXPERIENCE\nLed migration.\n")
    assert out.count("## EXPERIENCE") == 1


@pytest.mark.fast
def test_annotate_inserts_header_for_work_experience() -> None:
    out = _annotate_sections("WORK EXPERIENCE\nLed migration.\n")
    assert "## WORK EXPERIENCE" in out


@pytest.mark.fast
def test_annotate_inserts_header_for_closed_list_match_case_insensitive() -> None:
    out = _annotate_sections("Work Experience\nFoo\n")
    assert "## Work Experience" in out


@pytest.mark.fast
def test_annotate_does_not_eat_year_line() -> None:
    out = _annotate_sections("2024\nFoo\n")
    assert "## 2024" not in out


@pytest.mark.fast
def test_annotate_does_not_eat_version_token() -> None:
    out_cpp = _annotate_sections("C++17\nFoo\n")
    out_py = _annotate_sections("Python 3.10\nFoo\n")
    assert "##" not in out_cpp
    assert "##" not in out_py


@pytest.mark.fast
def test_annotate_does_not_eat_short_acronyms() -> None:
    """Short skill/tooling tokens on their own line must not become section headers."""
    for token in ("AWS S3", "CI/CD", "R&D", "IT/OPS"):
        out = _annotate_sections(f"{token}\nFoo\n")
        assert f"## {token}" not in out, f"{token!r} should not be promoted to ## header"


@pytest.mark.fast
def test_annotate_does_not_double_annotate() -> None:
    out = _annotate_sections("## Experience\nExperience\nFoo\n")
    assert out.count("## Experience") == 1


@pytest.mark.fast
@pytest.mark.parametrize(
    "cz_header",
    [
        "pracovní zkušenosti",
        "zkušenosti",
        "vzdělání",
        "dovednosti",
        "nejčastější dovednosti",
        "jazyky",
        "certifikace",
        "ocenění",
        "publikace",
        "projekty",
        "shrnutí",
        "profil",
        "kontakt",
    ],
)
def test_section_vocabulary_cz(cz_header: str) -> None:
    """CZ section aliases get promoted, preserving original casing.

    Also covers casing variants (title-case) and diacritic-stripped variants
    (`Vzdelani` for the model that strips diacritics) since both should hash
    to the same normalised key.
    """
    title = cz_header.title()
    out = _annotate_sections(f"{title}\nFoo\n")
    assert f"## {title}" in out, f"{title!r} should be promoted to ## header"

    no_diacritics = "".join(
        c for c in unicodedata.normalize("NFD", title) if unicodedata.category(c) != "Mn"
    )
    if no_diacritics != title:
        out2 = _annotate_sections(f"{no_diacritics}\nFoo\n")
        assert f"## {no_diacritics}" in out2, (
            f"diacritic-stripped {no_diacritics!r} should also be promoted"
        )


@pytest.mark.fast
@pytest.mark.parametrize(
    "en_header",
    [
        "Professional Experience",
        "Languages",
        "Certifications",
        "Honors-Awards",
        "Awards",
        "Publications",
        "Contact",
    ],
)
def test_section_vocabulary_en_extended(en_header: str) -> None:
    """New title-case EN aliases get promoted."""
    out = _annotate_sections(f"{en_header}\nFoo\n")
    assert f"## {en_header}" in out, f"{en_header!r} should be promoted to ## header"


@pytest.mark.fast
def test_long_all_caps_czech_header_is_promoted() -> None:
    """Regression: a long all-caps Czech header (> 40 chars, with diacritics)
    must still be recognised as a section header.

    Before this fix the `_MAX_HEADER_CHARS=40` gate ran first and the ASCII-only
    all-caps regex rejected diacritic uppercase letters — both combined hid
    real Czech CV headers like the one below from the section annotator,
    shifting section boundaries and breaking downstream anchor verification.
    """
    header = "PRACOVNÍ ZKUŠENOSTI A PROFESNÍ KARIÉRA V BANCE"
    assert len(header) > 40
    out = _annotate_sections(f"{header}\nVedl tým 4 datových vědců.\n")
    assert f"## {header}" in out, f"long all-caps CZ header should be promoted: {out!r}"


@pytest.mark.fast
def test_inline_languages_not_promoted() -> None:
    """A paragraph sentence that mentions 'languages' is not a section header."""
    sentence = "He learned several languages including French and German."
    out = _annotate_sections(f"{sentence}\nFoo\n")
    assert "##" not in out, f"sentence-length line should not be promoted: {out!r}"


@pytest.mark.fast
def test_inline_languages_list_not_promoted() -> None:
    """A long inline `Languages: ...` summary line is not promoted (length gate).

    Marek fixture has `Languages: Czech (native), English (C2), Russian (B1).`
    which we want to remain content, not become a header.
    """
    line = "Languages: Czech (native), English (C2), Russian (B1)."
    out = _annotate_sections(f"{line}\nFoo\n")
    assert "## Languages" not in out, f"inline languages-summary should not be promoted: {out!r}"


@pytest.mark.fast
async def test_observability_emits_start_and_done_for_docx_fixture() -> None:
    fixture = _FIXTURE_DIR / "01_junior_da_novotny.docx"
    data = fixture.read_bytes()
    # Fail loudly on unresolved LFS pointers (fresh checkout w/o `git lfs pull`
    # or CI without `lfs: true`). Otherwise the test would surface as a cryptic
    # zip-parse failure inside extract_text. — Copilot PR #2 finding.
    if data.startswith(b"version https://git-lfs.github.com/"):
        pytest.fail(
            f"{fixture.name} is an unresolved LFS pointer. "
            "Run `git lfs pull` (CI uses `actions/checkout@v4` with `lfs: true`)."
        )
    events: list[dict[str, Any]] = []
    with subscribe(events.append):
        result = await extract_text(data, fixture.name)
    assert isinstance(result, str)

    starts = [e for e in events if e["event"] == "start" and e["stage"] == "ingest"]
    dones = [e for e in events if e["event"] == "done" and e["stage"] == "ingest"]
    assert starts, f"expected a start event, got {events!r}"
    assert dones, f"expected a done event, got {events!r}"

    done = dones[0]
    assert done["chars"] >= 200
    assert "duration_ms" in done
    assert isinstance(done["duration_ms"], int)
    assert done["duration_ms"] >= 0
    assert "duration_ms" not in starts[0]


@pytest.mark.fast
async def test_observability_emits_rejected_for_unknown_suffix() -> None:
    events: list[dict[str, Any]] = []
    with subscribe(events.append):
        result = await extract_text(b"hello", "notes.txt")
    assert isinstance(result, StageFailure)

    rejected = [
        e
        for e in events
        if e["event"] == "rejected"
        and e["stage"] == "ingest"
        and e.get("reason") == "unknown_suffix"
    ]
    assert rejected, f"expected rejected/unknown_suffix event, got {events!r}"
    assert "duration_ms" in rejected[0]
    assert isinstance(rejected[0]["duration_ms"], int)


_FIXTURE_DIR = Path(__file__).parent / "fixtures" / "cvs"


@pytest.mark.slow
def test_cv_fixtures_corpus_not_empty() -> None:
    """Guard: empty parametrize silently auto-skips. Make corpus regression loud."""
    pdfs = list(_FIXTURE_DIR.glob("*.pdf"))
    docxs = list(_FIXTURE_DIR.glob("*.docx"))
    if not pdfs and not docxs:
        pytest.fail("no CV fixtures found in tests/fixtures/cvs/ — corpus regression")


@pytest.mark.slow
@pytest.mark.parametrize(
    "fixture_path",
    sorted(list(_FIXTURE_DIR.glob("*.pdf")) + list(_FIXTURE_DIR.glob("*.docx"))),
    ids=lambda p: p.name,
)
async def test_real_fixtures_extract_minimum_chars(fixture_path: Path) -> None:
    result = await extract_text(fixture_path.read_bytes(), fixture_path.name)
    assert isinstance(result, str), f"expected str, got {result!r}"
    assert len(result) >= 200


@pytest.mark.slow
async def test_scanned_pdf_returns_scanned_failure() -> None:
    reportlab_canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    buf = BytesIO()
    c = reportlab_canvas.Canvas(buf)
    c.setFillColorRGB(0.2, 0.2, 0.2)
    c.rect(50, 50, 200, 200, fill=1, stroke=0)
    c.showPage()
    c.save()
    pdf_bytes = buf.getvalue()

    # Precondition: the synthesized PDF must actually be text-poor on both
    # parsers; otherwise the detector is being silently bypassed.
    import pdfplumber as _pdfplumber
    import pypdf as _pypdf

    reader = _pypdf.PdfReader(BytesIO(pdf_bytes))
    pypdf_chars = len("\n".join(p.extract_text() or "" for p in reader.pages).strip())
    with _pdfplumber.open(BytesIO(pdf_bytes)) as plumber_pdf:
        plumber_chars = len("\n".join(p.extract_text() or "" for p in plumber_pdf.pages).strip())
    if pypdf_chars >= 100 or plumber_chars >= 100:
        pytest.fail(
            "scanned-pdf fixture synthesis produced extractable text "
            f"(pypdf={pypdf_chars}, pdfplumber={plumber_chars}) — "
            "fix the synthesis or commit a real fixture"
        )

    result = await extract_text(pdf_bytes, "scanned.pdf")
    assert isinstance(result, StageFailure)
    assert result.stage == "ingest"
    assert result.user_message == SCANNED_MSG
