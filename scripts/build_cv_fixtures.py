"""Author the test CV fixtures.

T04 covers CV #1 (junior, DOCX) and CV #8 (senior, PDF). T06 extends this
script with CVs #2–7, #9, #10. PDFs use a deliberately messy two-column
reportlab layout so pdfplumber's column-aware fallback gets exercised.

Run: ``uv run python scripts/build_cv_fixtures.py``
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer,
)

FIXTURES = Path("tests/fixtures/cvs")

# Register DejaVu so Czech diacritics survive — reportlab's built-in Helvetica /
# Times-Roman lack glyphs for ř / Č / š / é. Path is the standard Debian/Ubuntu
# location; if absent, we fall back to the built-ins (and accept lossy output).
_DEJAVU_DIR = Path("/usr/share/fonts/truetype/dejavu")
_FONT_SANS = "Helvetica"
_FONT_SANS_BOLD = "Helvetica-Bold"
_FONT_SANS_OBLIQUE = "Helvetica-Oblique"
_FONT_SERIF = "Times-Roman"
_FONT_SERIF_ITALIC = "Times-Italic"
if (_DEJAVU_DIR / "DejaVuSans.ttf").exists():
    pdfmetrics.registerFont(TTFont("DejaVuSans", str(_DEJAVU_DIR / "DejaVuSans.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", str(_DEJAVU_DIR / "DejaVuSans-Bold.ttf")))
    pdfmetrics.registerFont(
        TTFont("DejaVuSans-Oblique", str(_DEJAVU_DIR / "DejaVuSans-Oblique.ttf"))
    )
    pdfmetrics.registerFont(TTFont("DejaVuSerif", str(_DEJAVU_DIR / "DejaVuSerif.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVuSerif-Bold", str(_DEJAVU_DIR / "DejaVuSerif-Bold.ttf")))
    pdfmetrics.registerFont(
        TTFont("DejaVuSerif-Italic", str(_DEJAVU_DIR / "DejaVuSerif-Italic.ttf"))
    )
    # Bind bold/italic variants so <b>/<i> inside Paragraph picks them up.
    pdfmetrics.registerFontFamily(
        "DejaVuSans",
        normal="DejaVuSans",
        bold="DejaVuSans-Bold",
        italic="DejaVuSans-Oblique",
        boldItalic="DejaVuSans-Bold",
    )
    pdfmetrics.registerFontFamily(
        "DejaVuSerif",
        normal="DejaVuSerif",
        bold="DejaVuSerif-Bold",
        italic="DejaVuSerif-Italic",
        boldItalic="DejaVuSerif-Bold",
    )
    _FONT_SANS = "DejaVuSans"
    _FONT_SANS_BOLD = "DejaVuSans-Bold"
    _FONT_SANS_OBLIQUE = "DejaVuSans-Oblique"
    _FONT_SERIF = "DejaVuSerif"
    _FONT_SERIF_ITALIC = "DejaVuSerif-Italic"


# ---------- CV #1 — Jan Novotný, Junior Data Analyst (DOCX) ----------


def build_junior_docx(out_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("Jan Novotný", level=0)
    h.alignment = 1  # center
    contact = doc.add_paragraph()
    contact.alignment = 1
    contact.add_run("Junior Data Analyst — Prague, Czech Republic\n")
    contact.add_run("jan.novotny@example.cz | +420 777 123 456 | linkedin.com/in/jannovotnycz")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Junior Data Analyst with 1 year of professional experience at Mall.cz, "
        "focused on retail analytics and reporting automation in SQL and Python. "
        "Comfortable with PostgreSQL, dbt 1.7, Looker and basic pandas workflows. "
        "Looking to grow into a Data Scientist role over the next two years."
    )

    doc.add_heading("Experience", level=1)
    p = doc.add_paragraph()
    p.add_run("Junior Data Analyst — Mall.cz, Prague").bold = True
    p.add_run("\nJune 2025 – present (1 year)")
    doc.add_paragraph(
        "Built daily revenue and margin dashboards in Looker covering 4 product "
        "categories, replacing a manual Excel process and reducing reporting "
        "turnaround from 2 days to 4 hours."
    )
    doc.add_paragraph(
        "Owned the dbt model layer for the marketing mart: 18 dbt 1.7 models on "
        "PostgreSQL 15, with column-level tests covering 92% of business-critical fields."
    )
    doc.add_paragraph(
        "Wrote ad-hoc Python (pandas 2.2) analyses for the merchandising team — "
        "weekly cohort retention, basket-size segmentation, and a one-off churn "
        "investigation that flagged a 6.4% drop in repeat purchases among the "
        "Home & Garden segment."
    )
    doc.add_paragraph(
        "Took over the on-call data-quality rotation in November 2025 (1 week per "
        "month); resolved 11 alerts during my first rotation without escalation."
    )

    doc.add_heading("Education", level=1)
    p = doc.add_paragraph()
    p.add_run("Bachelor of Economics and Management — VŠE Prague").bold = True
    p.add_run("\n2022 – 2025; thesis on revenue forecasting using Prophet")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(
        "SQL (PostgreSQL 15, BigQuery), Python 3.11 (pandas 2.2, numpy), dbt 1.7, "
        "Looker, Git, basic Airflow 2.9 DAG maintenance."
    )
    doc.add_paragraph("Languages: Czech (native), English (C1, FCE 2021).")

    doc.save(out_path)


# ---------- CV #8 — Tomáš Dvořák, Staff ML Engineer (two-column PDF) ----------


@dataclass
class CVBlock:
    heading: str
    paragraphs: list[str]


def _senior_blocks() -> list[CVBlock]:
    return [
        CVBlock(
            "Summary",
            [
                "Staff Machine Learning Engineer with 13 years of experience across "
                "consumer security (Avast), travel (Kiwi.com) and retail banking "
                "(ČSOB). Tech-lead for two ML platforms serving 40M+ daily inferences. "
                "Currently responsible for the ML inference platform at ČSOB Prague, "
                "leading a team of 6 engineers and 2 data scientists."
            ],
        ),
        CVBlock(
            "Experience",
            [
                "<b>Staff ML Engineer — ČSOB, Prague</b><br/>"
                "January 2023 – present<br/>"
                "Tech lead for the ML inference platform serving 12M+ daily scoring "
                "calls across fraud, credit-risk and customer-360 use cases. "
                "Migrated the legacy SAS scoring stack to a Python 3.11 / FastAPI "
                "service running on Kubernetes 1.29, cutting median inference "
                "latency from 240 ms to 38 ms and reducing infra cost by 41%.",
                "Founded the ML platform guild (8 engineers across 3 squads) to "
                "standardise feature stores and evaluation harnesses. Authored the "
                "internal RFC on shadow-deployment and online-eval that is now the "
                "default release path for credit-risk models.",
                "<b>Senior ML Engineer — Kiwi.com, Brno</b><br/>"
                "March 2018 – December 2022<br/>"
                "Owned the ranking pipeline for the flight-search product (28M "
                "monthly active users). Replaced an XGBoost 0.9 ranker with a "
                "two-tower TensorFlow 2.8 retrieval-then-ranking architecture, "
                "lifting click-through on the top-3 results by 17.4% on a 90-day "
                "A/B test (p &lt; 0.001).",
                "Mentored 4 junior ML engineers, ran the bi-weekly modelling "
                "review, and wrote the team's first eval playbook covering offline "
                "AUC parity, online holdout, and post-launch latency budgets.",
                "<b>Machine Learning Engineer — Avast, Prague</b><br/>"
                "September 2013 – February 2018<br/>"
                "Built the gradient-boosted malware classifier for the Avast Free "
                "Antivirus Windows engine, processing 4.2 billion file scans per "
                "month at peak. Reduced false-positive rate from 0.32% to 0.11% "
                "while holding recall flat, using a calibrated LightGBM 2.1 model "
                "with weekly retraining.",
            ],
        ),
        CVBlock(
            "Education",
            [
                "<b>Ing. (M.Sc.) in Computer Science — ČVUT FIT, Prague</b><br/>"
                "2011 – 2013, focus on machine learning and pattern recognition.",
                "<b>Bc. (B.Sc.) in Computer Science — VUT Brno</b><br/>2008 – 2011.",
            ],
        ),
        CVBlock(
            "Skills",
            [
                "Python 3.11, TensorFlow 2.8, PyTorch 2.3, LightGBM 2.1, "
                "scikit-learn 1.5, FastAPI, Kubernetes 1.29, Argo Workflows, "
                "MLflow 2.10, Feast 0.40, BigQuery, PostgreSQL, Kafka 3.7.",
                "Leadership: tech lead (6 engineers + 2 DS), guild founder, "
                "RFC author, cross-team incident commander.",
                "Languages: Czech (native), English (C2), Slovak (fluent).",
            ],
        ),
        CVBlock(
            "Selected projects",
            [
                "<b>Project Hermes (ČSOB, 2024)</b> — real-time fraud-scoring "
                "service, 38 ms p50, deployed across 3 EU regions; reduced manual "
                "review queue by 24%.",
                "<b>Two-tower ranker (Kiwi.com, 2021)</b> — see Experience; "
                "presented at MLPrague 2022.",
            ],
        ),
    ]


def _build_messy_pdf(out_path: Path, blocks: list[CVBlock]) -> None:
    """Two-column reportlab layout with deliberate stressors for pdfplumber.

    Stressors:
      * narrow left column + wider right column (uneven frames),
      * a header band the columns flow under,
      * mixed font families across blocks,
      * a footer band with page number + date stamp.
    """

    doc = BaseDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=2.4 * cm,
        bottomMargin=2.0 * cm,
        title="Tomáš Dvořák — CV",
        author="Tomáš Dvořák",
    )

    page_w, page_h = A4
    gutter = 0.6 * cm
    left_w = (page_w - 1.6 * cm - 1.6 * cm - gutter) * 0.42
    right_w = (page_w - 1.6 * cm - 1.6 * cm - gutter) * 0.58
    body_top = page_h - 2.4 * cm - 1.6 * cm  # under the header band
    body_h = body_top - 2.0 * cm

    left_frame = Frame(1.6 * cm, 2.0 * cm, left_w, body_h, id="left", showBoundary=0)
    right_frame = Frame(
        1.6 * cm + left_w + gutter,
        2.0 * cm,
        right_w,
        body_h,
        id="right",
        showBoundary=0,
    )

    def _draw_chrome(canvas, _doc) -> None:
        canvas.saveState()
        # Header band — name + role
        canvas.setFont(_FONT_SANS_BOLD, 18)
        canvas.drawString(1.6 * cm, page_h - 1.8 * cm, "Tomáš Dvořák")
        canvas.setFont(_FONT_SANS, 11)
        canvas.drawString(
            1.6 * cm,
            page_h - 2.4 * cm,
            "Staff Machine Learning Engineer · Prague, Czech Republic",
        )
        canvas.setFont(_FONT_SANS_OBLIQUE, 9)
        canvas.drawString(
            1.6 * cm,
            page_h - 2.9 * cm,
            "tomas.dvorak@example.cz · +420 602 555 100 · linkedin.com/in/tomasdvorakml",
        )
        # Footer band — page number + version stamp
        canvas.setFont(_FONT_SERIF_ITALIC, 8)
        canvas.drawString(1.6 * cm, 1.2 * cm, "CV — last updated April 2026")
        canvas.drawRightString(page_w - 1.6 * cm, 1.2 * cm, f"Page {_doc.page}")
        canvas.restoreState()

    template = PageTemplate(id="two-col", frames=[left_frame, right_frame], onPage=_draw_chrome)
    doc.addPageTemplates([template])

    base = getSampleStyleSheet()
    h_style = ParagraphStyle(
        "Heading",
        parent=base["Heading2"],
        fontName=_FONT_SANS_BOLD,
        fontSize=12,
        spaceBefore=8,
        spaceAfter=4,
        textColor="#1a3a6b",
    )
    body_style = ParagraphStyle(
        "Body",
        parent=base["BodyText"],
        fontName=_FONT_SERIF,
        fontSize=10,
        leading=13,
        spaceAfter=6,
    )

    flow: list = []
    for block in blocks:
        flow.append(Paragraph(block.heading, h_style))
        for para in block.paragraphs:
            flow.append(Paragraph(para, body_style))
        flow.append(Spacer(1, 4))

    doc.build(flow)


def build_senior_pdf(out_path: Path) -> None:
    _build_messy_pdf(out_path, _senior_blocks())


# ---------- Golden .txt extraction ----------


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def main() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)

    junior_docx = FIXTURES / "01_junior_da_novotny.docx"
    senior_pdf = FIXTURES / "08_staff_ml_engineer_dvorak.pdf"
    junior_txt = FIXTURES / "01_junior_da_novotny.txt"
    senior_txt = FIXTURES / "08_staff_ml_engineer_dvorak.txt"

    build_junior_docx(junior_docx)
    build_senior_pdf(senior_pdf)

    junior_txt.write_text(extract_docx_text(junior_docx) + "\n", encoding="utf-8")
    senior_txt.write_text(extract_pdf_text(senior_pdf) + "\n", encoding="utf-8")

    print(f"wrote {junior_docx}")
    print(f"wrote {senior_pdf}")
    print(f"wrote {junior_txt}")
    print(f"wrote {senior_txt}")


if __name__ == "__main__":
    main()
